"""
Conversion Service – Job-Oriented Architecture
Handles PDF → Word, Text, HTML, Images
Crash-hardened edition – 2025-09
"""
import logging
import os
import re
import uuid
import io
from pathlib import Path
from typing import Any, Dict, List, Optional, BinaryIO

from src.services import FileManagementService
from src.jobs import JobOperationsWrapper, JobOperations
from src.models import JobStatus

logger = logging.getLogger(__name__)

# -------------- graceful library loading ------------------------------------
try:
    import fitz  # PyMuPDF
    PDF_LIBS_AVAILABLE = True
except ImportError:  # pragma: no cover
    PDF_LIBS_AVAILABLE = False
    logging.warning("PDF libs unavailable – install fitz")

try:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:  # pragma: no cover
    DOCX_AVAILABLE = False
    logging.warning("python-docx unavailable – DOCX export disabled")

try:
    import openpyxl
    from openpyxl.utils import get_column_letter
    OPENPYXL_AVAILABLE = True
except ImportError:  # pragma: no cover
    OPENPYXL_AVAILABLE = False
    logging.warning("openpyxl unavailable – Excel export disabled")

try:
    from PIL import Image, ImageDraw, ImageFont
    PIL_AVAILABLE = True
except ImportError:  # pragma: no cover
    PIL_AVAILABLE = False
    logging.warning("Pillow unavailable – image conversion disabled")

# ------------------------------------------------------------------------------
class ConversionService:
    """Convert PDFs to docx / txt / html / images – crash-hardened."""
    
    def __init__(self, file_service: Optional[FileManagementService] = None):
        self.file_service = file_service or FileManagementService()
        self.supported_formats = ("docx", "txt", "html", "images", "xlsx")

    # --------------------------------------------------------------------------
    # PUBLIC ENTRY POINTS
    # --------------------------------------------------------------------------
    
    def convert_pdf_data(self, file_data: bytes, target_format: str, 
                        options: Optional[Dict[str, Any]] = None,
                        original_filename: Optional[str] = None) -> Dict[str, Any]:
        """Convert *in-memory* PDF → target format. Always returns the same keys."""
        options = options or {}
        
        if target_format not in self.supported_formats:
            raise ValueError(f"Unsupported format {target_format}")

        if not PDF_LIBS_AVAILABLE:
            raise RuntimeError("PDF libraries not installed")

        temp_pdf_id: Optional[str] = None
        try:
            # Save file using FileManagementService
            temp_pdf_id, temp_pdf_path = self.file_service.save_file(file_data, original_filename or "temp.pdf")
            pdf_content = self._extract_pdf_content(Path(temp_pdf_path))

            converter = {
                "docx": self._convert_to_docx,
                "txt": self._convert_to_txt,
                "html": self._convert_to_html,
                "images": lambda content, opts: self._convert_to_images_with_pdf(temp_pdf_path, content, opts),
                "xlsx": self._convert_to_xlsx,
            }

            _method = converter[target_format]
            payload = _method(content=pdf_content, opts=options)
            
            # Guarantee keys that callers / frontend always read
            payload.setdefault("success", True)
            payload.setdefault("format", target_format)
            payload.setdefault("quality", options.get("quality", "medium"))
            payload.setdefault("original_filename", original_filename)
            payload.setdefault("original_size", len(file_data))
            
            return payload

        except Exception as exc:
            logger.exception("Conversion failed")
            # Always return same shape – UI will not blow up
            return {
                "success": False,
                "error": str(exc),
                "format": target_format,
                "quality": options.get("quality", "medium"),
                "original_filename": original_filename,
                "original_size": len(file_data),
            }
        finally:
            # Clean up temporary PDF file
            if temp_pdf_id:
                self.file_service.delete_file(self.file_service.get_file_path(temp_pdf_id))

    def get_conversion_preview(self, file_data: bytes, target_format: str, 
                              options: Optional[Dict[str, Any]] = None) -> Dict[str, Any]:
        """Light-weight preview; never crashes."""
        options = options or {}
        
        if not PDF_LIBS_AVAILABLE:
            return self._preview_fallback(file_data, target_format, "PDF libs missing")

        temp_pdf_id: Optional[str] = None
        try:
            # Save file using FileManagementService
            temp_pdf_id, temp_pdf_path = self.file_service.save_file(file_data, "preview.pdf")
            pdf_content = self._extract_pdf_content(Path(temp_pdf_path))
            
            return {
                "success": True,
                "original_size": len(file_data),
                "page_count": pdf_content.get("page_count", 0),
                "estimated_size": self._estimate_output_size(pdf_content, target_format),
                "estimated_time": self._estimate_conversion_time(pdf_content, target_format),
                "complexity": self._assess_complexity(pdf_content, target_format),
                "recommendations": self._get_recommendations(pdf_content, target_format, options),
                "supported_formats": self.supported_formats,
            }
        except Exception as exc:
            logger.exception("Preview failed")
            return self._preview_fallback(file_data, target_format, str(exc))
        finally:
            # Clean up temporary PDF file
            if temp_pdf_id:
                self.file_service.delete_file(self.file_service.get_file_path(temp_pdf_id))

    # --------------------------------------------------------------------------
    # INTERNALS – extraction
    # --------------------------------------------------------------------------
    
    def _extract_pdf_content(self, pdf_path: Path) -> Dict[str, Any]:
        """Return unified dict with pages, tables, images, text, meta."""
        try:
            with fitz.open(str(pdf_path)) as doc:
                if doc.is_closed or doc.needs_pass:
                    raise ValueError("PDF is corrupted or encrypted")
                
                content = {
                    "pages": [],
                    "tables": [],
                    "images": [],
                    "text": "",
                    "page_count": len(doc),
                    "metadata": doc.metadata or {},
                }
                
                for page_num, page in enumerate(doc, start=1):
                    text = page.get_text() or ""
                    content["text"] += text + "\n"
                    
                    tables = self._extract_tables_from_page(page)
                    images = self._extract_images_from_page(page)
                    
                    content["tables"].extend(tables)
                    content["images"].extend(images)
                    content["pages"].append({
                        "page_num": page_num,
                        "text": text,
                        "tables": tables,
                        "images": images,
                        "width": page.rect.width,
                        "height": page.rect.height,
                    })
                
                return content
                
        except Exception as e:
            logger.error(f"Failed to extract PDF content: {e}")
            raise ValueError(f"Failed to process PDF: {e}")

    # --------------------------------------------------------------------------
    # TABLE / IMAGE EXTRACTION
    # --------------------------------------------------------------------------
    
    def _extract_tables_from_page(self, page: fitz.Page) -> List[List[List[str]]]:
        """Extract tables from PDF page."""
        try:
            blocks = page.get_text("dict")["blocks"]
            tables: List[List[List[str]]] = []
            
            for block in blocks:
                if block.get("lines") and self._looks_like_table(block):
                    table = self._extract_table_data(block)
                    if table:
                        tables.append(table)
            
            return tables
        except Exception:
            return []

    @staticmethod
    def _looks_like_table(block: Dict[str, Any]) -> bool:
        """Determine if a text block looks like a table."""
        lines = block.get("lines", [])
        if len(lines) < 2:
            return False
        
        span_counts = [len(ln.get("spans", [])) for ln in lines]
        return len(set(span_counts)) <= 2 and max(span_counts) > 1

    @staticmethod
    def _extract_table_data(block: Dict[str, Any]) -> Optional[List[List[str]]]:
        """Extract table data from a block."""
        try:
            table: List[List[str]] = []
            for line in block["lines"]:
                row = [span["text"].strip() for span in line["spans"] if span["text"].strip()]
                if row:
                    table.append(row)
            return table if table else None
        except Exception:
            return None

    @staticmethod
    def _extract_images_from_page(page: fitz.Page) -> List[Dict[str, Any]]:
        """Return image metadata from PDF page."""
        try:
            return [
                {
                    "index": idx,
                    "xref": img[0],
                    "width": img[2],
                    "height": img[3],
                    "colorspace": img[4],
                    "bpc": img[5],
                }
                for idx, img in enumerate(page.get_images())
            ]
        except Exception:
            return []

    # --------------------------------------------------------------------------
    # CONVERTERS
    # --------------------------------------------------------------------------
    
    def _convert_to_docx(self, content: Dict[str, Any], opts: Dict[str, Any]) -> Dict[str, Any]:
        """Convert PDF content to DOCX format."""
        if not DOCX_AVAILABLE:
            raise RuntimeError("python-docx not installed")
        
        doc = Document()
        
        # Add title if available
        if content.get("metadata", {}).get("title"):
            title = doc.add_heading(content["metadata"]["title"], 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Process each page
        for page in content["pages"]:
            if page["page_num"] > 1:
                doc.add_page_break()
            
            # Add text content
            if page["text"].strip():
                for paragraph in page["text"].split("\n\n"):
                    if paragraph.strip():
                        doc.add_paragraph(paragraph.strip())
            
            # Add tables
            for table_data in page["tables"]:
                if not table_data:
                    continue
                
                max_cols = max(len(row) for row in table_data) or 1
                table = doc.add_table(rows=len(table_data), cols=max_cols)
                table.style = "Table Grid"
                
                for row_idx, row_data in enumerate(table_data):
                    for col_idx, cell_data in enumerate(row_data):
                        if col_idx < max_cols:
                            table.cell(row_idx, col_idx).text = str(cell_data)

        # Save to in-memory buffer
        docx_buffer = io.BytesIO()
        doc.save(docx_buffer)
        docx_data = docx_buffer.getvalue()
        docx_buffer.close()

        filename = f"converted_{self._secure_filename(content.get('metadata',{}).get('title') or 'document')}.docx"
        file_id, file_path = self.file_service.save_file(docx_data, filename)
        
        return {
            "success": True,
            "output_path": file_path,
            "filename": filename,
            "mime_type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            "file_size": self.file_service.get_file_size(file_path),
        }

    def _convert_to_xlsx(self, content: Dict[str, Any], opts: Dict[str, Any]) -> Dict[str, Any]:
        """Convert PDF content to Excel format."""
        if not OPENPYXL_AVAILABLE:
            raise RuntimeError("openpyxl not installed – Excel export unavailable")

        wb = openpyxl.Workbook()
        wb.remove(wb.active)  # Remove default sheet

        file_prefix = self._secure_filename(content.get("metadata", {}).get("title") or "document")
        any_table = False

        # Process each page with tables
        for page in content["pages"]:
            if not page["tables"]:
                continue
                
            any_table = True
            ws = wb.create_sheet(title=f"Page_{page['page_num']}")
            
            for table_idx, table_data in enumerate(page["tables"], start=1):
                start_row = ws.max_row + (2 if ws.max_row > 1 else 0)
                
                for row_idx, row_data in enumerate(table_data, start=start_row):
                    for col_idx, cell_val in enumerate(row_data, start=1):
                        ws.cell(row=row_idx, column=col_idx, value=str(cell_val))

                # Auto-size columns
                for col_idx in range(1, len(table_data[0]) + 1):
                    ws.column_dimensions[get_column_letter(col_idx)].auto_size = True

        if not any_table:
            raise ValueError("No tables found in PDF – nothing to export to Excel")

        # Save to in-memory buffer
        xlsx_buffer = io.BytesIO()
        wb.save(xlsx_buffer)
        xlsx_data = xlsx_buffer.getvalue()
        xlsx_buffer.close()

        filename = f"converted_{file_prefix}.xlsx"
        file_id, file_path = self.file_service.save_file(xlsx_data, filename)
        
        return {
            "success": True,
            "output_path": file_path,
            "filename": filename,
            "mime_type": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            "file_size": self.file_service.get_file_size(file_path),
        }

    def _convert_to_txt(self, content: Dict[str, Any], opts: Dict[str, Any]) -> Dict[str, Any]:
        """Convert PDF content to text format."""
        quality = opts.get("quality", "medium")
        text = content["text"]
        
        if quality != "high":
            text = re.sub(r"\s+", " ", text)
            text = re.sub(r"\n\s*\n", "\n\n", text)
            
        filename = f"converted_{self._secure_filename(content.get('metadata',{}).get('title') or 'document')}.txt"
        text_data = text.encode('utf-8')
        
        file_id, file_path = self.file_service.save_file(text_data, filename)
        
        return {
            "success": True,
            "output_path": file_path,
            "filename": filename,
            "mime_type": "text/plain",
            "file_size": self.file_service.get_file_size(file_path),
        }

    def _convert_to_html(self, content: Dict[str, Any], opts: Dict[str, Any]) -> Dict[str, Any]:
        """Convert PDF content to HTML format."""
        html = self._generate_html_content(content, opts)
        filename = f"converted_{self._secure_filename(content.get('metadata',{}).get('title') or 'document')}.html"
        html_data = html.encode('utf-8')
        
        file_id, file_path = self.file_service.save_file(html_data, filename)
        
        return {
            "success": True,
            "output_path": file_path,
            "filename": filename,
            "mime_type": "text/html",
            "file_size": self.file_service.get_file_size(file_path),
        }

    def _convert_to_images_with_pdf(self, pdf_path: str, content: Dict[str, Any], opts: Dict[str, Any]) -> Dict[str, Any]:
        """PDF → PNG (one file per page) using the actual PDF file."""
        if not PIL_AVAILABLE:
            raise RuntimeError("Pillow not installed – image conversion unavailable")

        dpi = {"low": 72, "medium": 150, "high": 300}.get(opts.get("quality", "medium"), 150)
        image_files = []
        total_size = 0

        try:
            # Open the PDF document
            with fitz.open(pdf_path) as doc:
                if doc.is_closed or doc.needs_pass:
                    raise ValueError("PDF is corrupted or encrypted")

                for page_num in range(len(doc)):
                    page = doc.load_page(page_num)
                    
                    # Get pixmap with specified DPI
                    with page.get_pixmap(dpi=dpi) as pix:
                        # Convert pixmap to PNG bytes
                        image_data = pix.tobytes("png")
                        
                        filename = f"converted_page_{page_num + 1}.png"
                        
                        # Save through file service
                        file_id, file_path = self.file_service.save_file(image_data, filename)
                        image_files.append(file_path)
                        total_size += len(image_data)

                if not image_files:
                    raise RuntimeError("No images could be created")

                return {
                    "success": True,
                    "output_path": str(Path(image_files[0]).parent) if image_files else "",
                    "filename": "pages_converted.zip",
                    "mime_type": "application/zip",
                    "file_size": total_size,
                    "image_files": image_files,
                    "page_count": len(image_files)
                }

        except Exception as e:
            logger.error(f"Image conversion failed: {e}")
            # Fall back to the text-based approach
            return self._convert_to_images(content, opts)

    def _convert_to_images(self, content: Dict[str, Any], opts: Dict[str, Any]) -> Dict[str, Any]:
        """Fallback PDF → PNG using text rendering."""
        if not PIL_AVAILABLE:
            raise RuntimeError("Pillow not installed – image conversion unavailable")

        dpi = {"low": 72, "medium": 150, "high": 300}.get(opts.get("quality", "medium"), 150)
        image_files = []
        total_size = 0

        # Create individual PNG files for each page using text content
        for page_data in content["pages"]:
            page_num = page_data["page_num"]
            filename = f"converted_page_{page_num}.png"

            try:
                # Create a simple text-based image representation
                img_width, img_height = 800, 1000
                img = Image.new('RGB', (img_width, img_height), color='white')
                draw = ImageDraw.Draw(img)

                # Try to use a basic font
                try:
                    font = ImageFont.load_default()
                except:
                    font = None

                # Draw the page text
                text = page_data.get("text", f"Page {page_num}")[:1000]  # Limit text length
                if text.strip():
                    # Simple text wrapping
                    lines = []
                    words = text.split()
                    current_line = ""

                    for word in words[:100]:  # Limit words to prevent overflow
                        if len(current_line + word) < 80:  # Rough character limit per line
                            current_line += word + " "
                        else:
                            if current_line:
                                lines.append(current_line.strip())
                            current_line = word + " "

                    if current_line:
                        lines.append(current_line.strip())

                    # Draw lines
                    y_position = 50
                    for line in lines[:40]:  # Limit lines to fit on image
                        draw.text((50, y_position), line, fill='black', font=font)
                        y_position += 25

                # Save image to bytes
                img_buffer = io.BytesIO()
                img.save(img_buffer, format='PNG', dpi=(dpi, dpi))
                image_data = img_buffer.getvalue()
                img_buffer.close()

                file_id, file_path = self.file_service.save_file(image_data, filename)
                image_files.append(file_path)
                total_size += len(image_data)

            except Exception as e:
                logger.warning(f"Failed to create image for page {page_num}: {e}")
                # Create minimal placeholder
                placeholder_text = f"Page {page_num} - Image conversion failed"
                image_data = placeholder_text.encode('utf-8')
                file_id, file_path = self.file_service.save_file(image_data, f"page_{page_num}_error.txt")
                image_files.append(file_path)
                total_size += len(image_data)

        if not image_files:
            raise RuntimeError("No images could be created")

        return {
            "success": True,
            "output_path": str(Path(image_files[0]).parent) if image_files else "",
            "filename": "pages_converted.zip",
            "mime_type": "application/zip",
            "file_size": total_size,
            "image_files": image_files,
            "page_count": len(image_files)
        }

    # --------------------------------------------------------------------------
    # HTML generator
    # --------------------------------------------------------------------------
    
    def _generate_html_content(self, content: Dict[str, Any], opts: Dict[str, Any]) -> str:
        """Generate HTML content from PDF data."""
        css = ""
        if opts.get("include_css", True):
            css = """
            <style>
                body{font-family:Arial,sans-serif;margin:20px}
                .page{margin-bottom:30px}
                table{border-collapse:collapse;width:100%;margin:15px 0}
                th,td{border:1px solid #ddd;padding:8px;text-align:left}
                th{background:#f2f2f2}
            </style>
            """
            
        parts = [
            "<!DOCTYPE html><html><head>",
            "<meta charset='utf-8'><meta name='viewport' content='width=device-width,initial-scale=1'>",
            f"<title>Converted PDF</title>{css}</head><body>",
        ]
        
        if content.get("metadata", {}).get("title"):
            parts.append(f"<h1>{content['metadata']['title']}</h1>")
            
        for page in content["pages"]:
            parts.append(f'<div class="page"><h3>Page {page["page_num"]}</h3>')
            
            if page["text"].strip():
                for para in page["text"].split("\n\n"):
                    if para.strip():
                        parts.append(f"<p>{para.strip()}</p>")
                        
            for tbl in page["tables"]:
                if not tbl:
                    continue
                    
                parts.append("<table>")
                for r_idx, row in enumerate(tbl):
                    tag = "th" if r_idx == 0 else "td"
                    parts.append("<tr>")
                    for cell in row:
                        parts.append(f"<{tag}>{cell}</{tag}>")
                    parts.append("</tr>")
                parts.append("</table>")
                
            parts.append("</div>")
            
        parts.extend(["</body></html>"])
        return "".join(parts)

    # --------------------------------------------------------------------------
    # preview helpers
    # --------------------------------------------------------------------------
    
    def _preview_fallback(self, file_data: bytes, fmt: str, err: str) -> Dict[str, Any]:
        """Fallback preview when conversion fails."""
        return {
            "success": False,
            "error": err,
            "original_size": len(file_data),
            "page_count": 0,
            "estimated_size": 0,
            "estimated_time": 0,
            "complexity": "unknown",
            "recommendations": [],
            "supported_formats": self.supported_formats,
        }

    @staticmethod
    def _estimate_output_size(content: Dict[str, Any], fmt: str) -> int:
        """Estimate output file size."""
        base = len(content.get("text", ""))
        mul = {"docx": 1.2, "txt": 0.3, "html": 1.5, "images": 2.0}.get(fmt, 1.0)
        return int(base * mul)

    @staticmethod
    def _estimate_conversion_time(content: Dict[str, Any], fmt: str) -> int:
        """Estimate conversion time in seconds."""
        pct = content.get("page_count", 1)
        tpp = {"docx": 2, "txt": 0.5, "html": 1, "images": 3}.get(fmt, 1)
        return pct * tpp

    def _assess_complexity(self, content: Dict[str, Any], _: str) -> str:
        """Assess conversion complexity."""
        pc, tc, ic = content.get("page_count", 0), len(content["tables"]), len(content["images"])
        if pc > 50 or tc > 20 or ic > 100:
            return "high"
        if pc > 20 or tc > 10 or ic > 50:
            return "medium"
        return "low"

    @staticmethod
    def _get_recommendations(content: Dict[str, Any], fmt: str, _: Dict[str, Any]) -> List[str]:
        """Get conversion recommendations."""
        rec = []
        pc, tc, ic = content.get("page_count", 0), len(content["tables"]), len(content["images"])
        
        if pc > 100:
            rec.append("Large document – consider splitting.")
        if tc and fmt == "txt":
            rec.append("Tables found – DOCX preserves them better.")
        if ic and fmt == "txt":
            rec.append("Images found – TXT will lose them.")
            
        return rec

    @staticmethod
    def _secure_filename(name: str) -> str:
        """Sanitize filename to be filesystem-safe."""
        if not name:
            return "file"
        name = re.sub(r"[^\w\s-]", "", name)
        name = re.sub(r"[-\s]+", "-", name)
        return name.strip("-") or "file"

    # --------------------------------------------------------------------------
    # JOB PROCESSING METHODS
    # --------------------------------------------------------------------------
    
    def process_conversion_job(self, job_id: str, file_data: bytes, target_format: str, 
                              options: Dict[str, Any] = None) -> Dict[str, Any]:
        """Process a conversion job with provided file data and job status management."""
        try:
            # Get job information
            job = JobOperations.get_job(job_id=job_id)
            if not job:
                raise ValueError(f"Job {job_id} not found")

            # Update job status to processing
            JobOperationsWrapper.update_job_status_safely(
                job_id=job.job_id,
                status=JobStatus.PROCESSING
            )

            # Get job details
            input_data = job.input_data or {}
            original_filename = input_data.get('original_filename')
            
            if not input_data:
                raise ValueError(f"Job {job_id} malformed data")

            # Process the file data
            result = self.convert_pdf_data(
                file_data=file_data,
                target_format=target_format,
                options=options or {},
                original_filename=original_filename
            )

            # Update job status based on result
            if result.get('success', False):
                JobOperationsWrapper.update_job_status_safely(
                    job_id=job_id,
                    status=JobStatus.COMPLETED,
                    result=result
                )
            else:
                JobOperationsWrapper.update_job_status_safely(
                    job_id=job_id,
                    status=JobStatus.FAILED,
                    error_message=result.get('error', 'Conversion failed')
                )

            return result

        except Exception as e:
            JobOperationsWrapper.update_job_status_safely(
                job_id=job_id,
                status=JobStatus.FAILED,
                error_message=str(e)
            )
            logger.error(f"Error processing conversion job {job_id}: {str(e)}")
            raise

    # --------------------------------------------------------------------------
    # SERVICE STATUS AND HEALTH CHECKS
    # --------------------------------------------------------------------------
    
    def get_service_status(self) -> Dict[str, Any]:
        """Get current status of the conversion service."""
        return {
            'service_name': 'ConversionService',
            'supported_formats': self.supported_formats,
            'libraries_available': {
                'pdf': PDF_LIBS_AVAILABLE,
                'docx': DOCX_AVAILABLE,
                'excel': OPENPYXL_AVAILABLE,
                'images': PIL_AVAILABLE
            },
            'file_service_available': hasattr(self, 'file_service') and self.file_service is not None,
            'timestamp': '2025-09-16T00:00:00Z'  # Placeholder, use actual timestamp in production
        }

    def health_check(self) -> Dict[str, Any]:
        """Perform health check of the conversion service."""
        health_status = {
            'status': 'healthy',
            'service': 'ConversionService',
            'checks': {},
            'errors': []
        }

        try:
            # Check library availability
            lib_check = {
                'pdf_libs': PDF_LIBS_AVAILABLE,
                'docx_libs': DOCX_AVAILABLE,
                'excel_libs': OPENPYXL_AVAILABLE,
                'image_libs': PIL_AVAILABLE
            }
            
            health_status['checks']['libraries'] = lib_check
            
            # Check if any critical libraries are missing
            missing_libs = [lib for lib, available in lib_check.items() if not available]
            if missing_libs:
                health_status['status'] = 'degraded'
                health_status['errors'].append(f"Missing libraries: {', '.join(missing_libs)}")
                
            # Check file service availability
            if not hasattr(self, 'file_service') or self.file_service is None:
                health_status['status'] = 'unhealthy'
                health_status['errors'].append("File management service not available")
                
        except Exception as e:
            health_status['status'] = 'unhealthy'
            health_status['errors'].append(f"Health check failed: {str(e)}")
            
        return health_status