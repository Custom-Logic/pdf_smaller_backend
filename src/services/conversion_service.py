"""
Conversion Service – Job-Oriented Architecture
Handles PDF → Word, Text, HTML, Images, Excel
Refactored to match compression service patterns
"""
import io
import json
import logging
import re
from pathlib import Path
from typing import Any, Dict, List, Optional


from src.models import JobStatus

logger = logging.getLogger(__name__)

# -------------- graceful library loading ------------------------------------
try:
    import fitz  # PyMuPDF
    PDF_LIBS_AVAILABLE = True
except ImportError:
    PDF_LIBS_AVAILABLE = False
    logging.warning("PDF libs unavailable – install fitz")

try:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logging.warning("python-docx unavailable – DOCX export disabled")

try:
    import openpyxl
    from openpyxl.utils import get_column_letter
    OPENPYXL_AVAILABLE = True
except ImportError:
    OPENPYXL_AVAILABLE = False
    logging.warning("openpyxl unavailable – Excel export disabled")

try:
    from PIL import Image, ImageDraw, ImageFont
    PIL_AVAILABLE = True
except ImportError:
    PIL_AVAILABLE = False
    logging.warning("Pillow unavailable – image conversion disabled")


class ConversionService:
    """Convert PDFs to docx/txt/html/images/xlsx – follows compression service pattern."""
    
    def __init__(self, file_service = None):
        from src.services import ServiceRegistry
        self.file_service = file_service or ServiceRegistry.get_file_management_service()
        self.supported_formats = ("docx", "txt", "html", "images", "xlsx")

    # --------------------------------------------------------------------------
    # MAIN ENTRY POINT - Matches compression service pattern
    # --------------------------------------------------------------------------
    def process_conversion_job(self, job_id: str, file_data: bytes, target_format: str, 
                              options: Optional[Dict[str, Any]] = None) -> Dict[str, Any]:
        """Process conversion job with job status management - main entry point."""
        try:
            from src.main import job_operations_controller
            logger.debug(f"Starting conversion job {job_id} to {target_format}")
            
            # Verify job exists (it should have been created in the route/task)
            job = job_operations_controller.job_operations.get_job(job_id=job_id)
            if not job:
                logger.error(f"Job {job_id} not found")
                raise ValueError(f"Job {job_id} not found")

            # Update status to processing
            job_operations_controller.update_job_status_safely(
                job_id=job_id,
                status=JobStatus.PROCESSING,
                progress=10.0
            )
            logger.debug(f"Job {job_id} marked as processing")

            # Get original filename from job input data
            # src/services/conversion_service.py  (top of process_conversion_job)

            try:
                input_data = json.loads(job.input_data) if isinstance(job.input_data, str) else (job.input_data or {})
                original_filename = input_data.get('original_filename') or 'document.pdf'
            except Exception:
                logger.exception('Could not parse job.input_data, using defaults')
                input_data = {}
                original_filename = 'document.pdf'


            # Process the conversion
            result = self.convert_pdf_data(
                file_data=file_data,
                target_format=target_format,
                options=options or {},
                original_filename=original_filename
            )

            # Update job status based on result
            if result.get('success', False):
                job_operations_controller.update_job_status_safely(
                    job_id=job_id,
                    status=JobStatus.COMPLETED,
                    result=result,
                    progress=100.0
                )
                logger.info(f"Conversion job {job_id} completed successfully")
            else:
                error_msg = result.get('error', 'Unknown conversion error')
                job_operations_controller.update_job_status_safely(
                    job_id=job_id,
                    status=JobStatus.FAILED,
                    error_message=error_msg
                )
                logger.error(f"Conversion job {job_id} failed: {error_msg}")

            return result

        except Exception as e:
            logger.exception(f"Error processing conversion job {job_id}")
            job_operations_controller.update_job_status_safely(
                job_id=job_id,
                status=JobStatus.FAILED,
                error_message=str(e)
            )
            return {
                "success": False,
                "error": str(e),
                "format": target_format,
                "original_filename": options.get('original_filename')
            }

    def convert_pdf_data(self, file_data: bytes, target_format: str, 
                        options: Optional[Dict[str, Any]] = None,
                        original_filename: Optional[str] = None) -> Dict[str, Any]:
        """Convert PDF to target format - internal method."""
        options = options or {}
        
        if target_format.casefold() not in [val.casefold() for val in self.supported_formats]:
            raise ValueError(f"Unsupported format {target_format}")

        if not PDF_LIBS_AVAILABLE:
            raise RuntimeError("PDF libraries not installed")

        temp_pdf_id: Optional[str] = None
        try:
            # Save temporary PDF file
            temp_pdf_id, temp_pdf_path = self.file_service.save_file(
                file_data, 
                original_filename or "temp.pdf"
            )
            
            # Extract content from PDF
            pdf_content = self._extract_pdf_content(Path(temp_pdf_path))

            # Select appropriate converter
            if target_format == "docx":
                result = self._convert_to_docx(pdf_content, options)
            elif target_format == "txt":
                result = self._convert_to_txt(pdf_content, options)
            elif target_format == "html":
                result = self._convert_to_html(pdf_content, options)
            elif target_format == "images":
                result = self._convert_to_images_with_pdf(temp_pdf_path, pdf_content, options)
            elif target_format == "xlsx":
                result = self._convert_to_xlsx(pdf_content, options)
            else:
                raise ValueError(f"Unsupported format: {target_format}")
            
            # Add standard metadata
            result.setdefault("success", True)
            result.setdefault("format", target_format)
            result.setdefault("quality", options.get("quality", "medium"))
            result.setdefault("original_filename", original_filename)
            result.setdefault("original_size", len(file_data))
            
            return result

        except Exception as exc:
            logger.exception("Conversion failed")
            return {
                "success": False,
                "error": str(exc),
                "format": target_format,
                "quality": options.get("quality", "medium"),
                "original_filename": original_filename,
                "original_size": len(file_data) if file_data else 0,
            }
        finally:
            # Clean up temporary PDF file
            if temp_pdf_id and self.file_service:
                try:
                    self.file_service.delete_file(self.file_service.get_file_path(temp_pdf_id))
                except Exception as e:
                    logger.warning(f"Failed to cleanup temp file: {e}")

    def get_conversion_preview(self, file_data: bytes, target_format: str, 
                              options: Optional[Dict[str, Any]] = None) -> Dict[str, Any]:
        """Generate preview for conversion."""
        options = options or {}
        
        if not PDF_LIBS_AVAILABLE:
            return self._preview_fallback(file_data, target_format, "PDF libs missing")

        temp_pdf_id: Optional[str] = None
        try:
            temp_pdf_id, temp_pdf_path = self.file_service.save_file(file_data, "preview.pdf")
            pdf_content = self._extract_pdf_content(Path(temp_pdf_path))
            
            return {
                "success": True,
                "original_size": len(file_data),
                "page_count": pdf_content.get("page_count", 0),
                "estimated_size": self._estimate_output_size(pdf_content, target_format),
                "estimated_time": self._estimate_conversion_time(pdf_content, target_format),
                "complexity": self._assess_complexity(pdf_content, target_format),
                "recommendations": self._get_recommendations(pdf_content, target_format, options),
                "supported_formats": self.supported_formats,
            }
        except Exception as exc:
            logger.exception("Preview failed")
            return self._preview_fallback(file_data, target_format, str(exc))
        finally:
            if temp_pdf_id and self.file_service:
                try:
                    self.file_service.delete_file(self.file_service.get_file_path(temp_pdf_id))
                except Exception as e:
                    logger.warning(f"Failed to cleanup temp file: {e}")

    # --------------------------------------------------------------------------
    # PDF CONTENT EXTRACTION
    # --------------------------------------------------------------------------
    def _extract_pdf_content(self, pdf_path: Path) -> Dict[str, Any]:
        """Extract content from PDF file."""
        try:
            with fitz.open(str(pdf_path)) as doc:
                if doc.is_closed or doc.needs_pass:
                    raise ValueError("PDF is corrupted or encrypted")
                
                content = {
                    "pages": [],
                    "tables": [],
                    "images": [],
                    "text": "",
                    "page_count": len(doc),
                    "metadata": doc.metadata or {},
                }
                
                for page_num, page in enumerate(doc, start=1):
                    text = page.get_text() or ""
                    content["text"] += text + "\n"
                    
                    tables = self._extract_tables_from_page(page)
                    images = self._extract_images_from_page(page)
                    
                    content["tables"].extend(tables)
                    content["images"].extend(images)
                    content["pages"].append({
                        "page_num": page_num,
                        "text": text,
                        "tables": tables,
                        "images": images,
                        "width": page.rect.width,
                        "height": page.rect.height,
                    })
                
                return content
                
        except Exception as e:
            logger.error(f"Failed to extract PDF content: {e}")
            raise ValueError(f"Failed to process PDF: {e}")

    def _extract_tables_from_page(self, page: Any) -> List[List[List[str]]]:
        """Extract tables from PDF page."""
        try:
            blocks = page.get_text("dict")["blocks"]
            tables: List[List[List[str]]] = []
            
            for block in blocks:
                if block.get("lines") and self._looks_like_table(block):
                    table = self._extract_table_data(block)
                    if table:
                        tables.append(table)
            
            return tables
        except Exception:
            return []

    @staticmethod
    def _looks_like_table(block: Dict[str, Any]) -> bool:
        """Check if block looks like a table."""
        lines = block.get("lines", [])
        if len(lines) < 2:
            return False
        
        span_counts = [len(ln.get("spans", [])) for ln in lines]
        return len(set(span_counts)) <= 2 and max(span_counts) > 1

    @staticmethod
    def _extract_table_data(block: Dict[str, Any]) -> Optional[List[List[str]]]:
        """Extract table data from block."""
        try:
            table: List[List[str]] = []
            for line in block["lines"]:
                row = [span["text"].strip() for span in line["spans"] if span["text"].strip()]
                if row:
                    table.append(row)
            return table if table else None
        except Exception:
            return None

    @staticmethod
    def _extract_images_from_page(page: Any) -> List[Dict[str, Any]]:
        """Extract image metadata from page."""
        try:
            return [
                {
                    "index": idx,
                    "xref": img[0],
                    "width": img[2],
                    "height": img[3],
                    "colorspace": img[4],
                    "bpc": img[5],
                }
                for idx, img in enumerate(page.get_images())
            ]
        except Exception:
            return []

    # --------------------------------------------------------------------------
    # CONVERTERS
    # --------------------------------------------------------------------------
    def _convert_to_docx(self, content: Dict[str, Any], opts: Dict[str, Any]) -> Dict[str, Any]:
        """Convert to DOCX format."""
        if not DOCX_AVAILABLE:
            raise RuntimeError("python-docx not installed")
        
        doc = Document()
        
        # Add title if available
        if content.get("metadata", {}).get("title"):
            title = doc.add_heading(content["metadata"]["title"], 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Process pages
        for page in content["pages"]:
            if page["page_num"] > 1:
                doc.add_page_break()
            
            # Add text
            if page["text"].strip():
                for paragraph in page["text"].split("\n\n"):
                    if paragraph.strip():
                        doc.add_paragraph(paragraph.strip())
            
            # Add tables
            for table_data in page["tables"]:
                if not table_data:
                    continue
                
                max_cols = max(len(row) for row in table_data) if table_data else 1
                table = doc.add_table(rows=len(table_data), cols=max_cols)
                table.style = "Table Grid"
                
                for row_idx, row_data in enumerate(table_data):
                    for col_idx, cell_data in enumerate(row_data):
                        if col_idx < max_cols:
                            table.cell(row_idx, col_idx).text = str(cell_data)

        # Save to bytes
        docx_buffer = io.BytesIO()
        doc.save(docx_buffer)
        docx_data = docx_buffer.getvalue()
        docx_buffer.close()

        # Save using file service
        filename = f"converted_{self._secure_filename(content.get('metadata',{}).get('title') or 'document')}.docx"
        file_id, file_path = self.file_service.save_file(docx_data, filename)
        
        return {
            "success": True,
            "output_path": file_path,
            "filename": filename,
            "mime_type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            "file_size": len(docx_data),
        }

    def _convert_to_xlsx(self, content: Dict[str, Any], opts: Dict[str, Any]) -> Dict[str, Any]:
        """Convert to Excel format."""
        if not OPENPYXL_AVAILABLE:
            raise RuntimeError("openpyxl not installed")

        wb = openpyxl.Workbook()
        wb.remove(wb.active)  # Remove default sheet

        file_prefix = self._secure_filename(content.get("metadata", {}).get("title") or "document")
        any_table = False

        # Process pages with tables
        for page in content["pages"]:
            if not page["tables"]:
                continue
                
            any_table = True
            ws = wb.create_sheet(title=f"Page_{page['page_num']}")
            
            current_row = 1
            for table_idx, table_data in enumerate(page["tables"], start=1):
                if table_idx > 1:
                    current_row += 2  # Add spacing between tables
                
                for row_data in table_data:
                    for col_idx, cell_val in enumerate(row_data, start=1):
                        ws.cell(row=current_row, column=col_idx, value=str(cell_val))
                    current_row += 1

                # Auto-size columns
                if table_data:
                    for col_idx in range(1, min(len(table_data[0]) + 1, 100)):
                        col_letter = get_column_letter(col_idx)
                        ws.column_dimensions[col_letter].auto_size = True

        if not any_table:
            # Create a sheet with text content if no tables
            ws = wb.create_sheet(title="Content")
            ws.cell(row=1, column=1, value=content.get("text", "No content extracted"))

        # Save to bytes
        xlsx_buffer = io.BytesIO()
        wb.save(xlsx_buffer)
        xlsx_data = xlsx_buffer.getvalue()
        xlsx_buffer.close()

        filename = f"converted_{file_prefix}.xlsx"
        file_id, file_path = self.file_service.save_file(xlsx_data, filename)
        
        return {
            "success": True,
            "output_path": file_path,
            "filename": filename,
            "mime_type": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            "file_size": len(xlsx_data),
        }

    def _convert_to_txt(self, content: Dict[str, Any], opts: Dict[str, Any]) -> Dict[str, Any]:
        """Convert to text format."""
        quality = opts.get("quality", "medium")
        text = content["text"]
        
        if quality != "high":
            text = re.sub(r"\s+", " ", text)
            text = re.sub(r"\n\s*\n", "\n\n", text)
            
        filename = f"converted_{self._secure_filename(content.get('metadata',{}).get('title') or 'document')}.txt"
        text_data = text.encode('utf-8')
        
        file_id, file_path = self.file_service.save_file(text_data, filename)
        
        return {
            "success": True,
            "output_path": file_path,
            "filename": filename,
            "mime_type": "text/plain",
            "file_size": len(text_data),
        }

    def _convert_to_html(self, content: Dict[str, Any], opts: Dict[str, Any]) -> Dict[str, Any]:
        """Convert to HTML format."""
        html = self._generate_html_content(content, opts)
        filename = f"converted_{self._secure_filename(content.get('metadata',{}).get('title') or 'document')}.html"
        html_data = html.encode('utf-8')
        
        file_id, file_path = self.file_service.save_file(html_data, filename)
        
        return {
            "success": True,
            "output_path": file_path,
            "filename": filename,
            "mime_type": "text/html",
            "file_size": len(html_data),
        }

    def _convert_to_images_with_pdf(self, pdf_path: str, content: Dict[str, Any], opts: Dict[str, Any]) -> Dict[str, Any]:
        """Convert PDF to images."""
        if not PDF_LIBS_AVAILABLE:
            raise RuntimeError("PDF libraries not installed")

        dpi = {"low": 72, "medium": 150, "high": 300}.get(opts.get("quality", "medium"), 150)
        image_files = []
        total_size = 0

        try:
            with fitz.open(pdf_path) as doc:
                if doc.is_closed or doc.needs_pass:
                    raise ValueError("PDF is corrupted or encrypted")

                for page_num in range(len(doc)):
                    page = doc.load_page(page_num)
                    
                    # Get pixmap with specified DPI
                    mat = fitz.Matrix(dpi/72.0, dpi/72.0)
                    pix = page.get_pixmap(matrix=mat)
                    
                    # Convert to PNG bytes
                    img_data = pix.tobytes("png")
                    
                    filename = f"page_{page_num + 1}.png"
                    file_id, file_path = self.file_service.save_file(img_data, filename)
                    image_files.append(file_path)
                    total_size += len(img_data)

                if not image_files:
                    raise RuntimeError("No images could be created")

                return {
                    "success": True,
                    "output_path": str(Path(image_files[0]).parent) if image_files else "",
                    "filename": "converted_images.zip",
                    "mime_type": "application/zip",
                    "file_size": total_size,
                    "image_files": image_files,
                    "page_count": len(image_files)
                }

        except Exception as e:
            logger.error(f"Image conversion failed: {e}")
            raise

    def _generate_html_content(self, content: Dict[str, Any], opts: Dict[str, Any]) -> str:
        """Generate HTML from content."""
        css = ""
        if opts.get("include_css", True):
            css = """
            <style>
                body{font-family:Arial,sans-serif;margin:20px;line-height:1.6;}
                .page{margin-bottom:30px;padding:20px;border-bottom:1px solid #ddd;}
                table{border-collapse:collapse;width:100%;margin:15px 0;}
                th,td{border:1px solid #ddd;padding:8px;text-align:left;}
                th{background:#f2f2f2;font-weight:bold;}
                h1{color:#333;}h3{color:#666;margin-top:20px;}
                p{margin:10px 0;}
            </style>
            """
            
        parts = [
            "<!DOCTYPE html><html><head>",
            "<meta charset='utf-8'><meta name='viewport' content='width=device-width,initial-scale=1'>",
            f"<title>{content.get('metadata', {}).get('title', 'Converted PDF')}</title>{css}</head><body>",
        ]
        
        if content.get("metadata", {}).get("title"):
            parts.append(f"<h1>{content['metadata']['title']}</h1>")
            
        for page in content["pages"]:
            parts.append(f'<div class="page"><h3>Page {page["page_num"]}</h3>')
            
            if page["text"].strip():
                for para in page["text"].split("\n\n"):
                    if para.strip():
                        parts.append(f"<p>{self._escape_html(para.strip())}</p>")
                        
            for tbl in page["tables"]:
                if not tbl:
                    continue
                    
                parts.append("<table>")
                for r_idx, row in enumerate(tbl):
                    tag = "th" if r_idx == 0 else "td"
                    parts.append("<tr>")
                    for cell in row:
                        parts.append(f"<{tag}>{self._escape_html(str(cell))}</{tag}>")
                    parts.append("</tr>")
                parts.append("</table>")
                
            parts.append("</div>")
            
        parts.extend(["</body></html>"])
        return "".join(parts)

    @staticmethod
    def _escape_html(text: str) -> str:
        """Escape HTML special characters."""
        return (text.replace("&", "&amp;")
                   .replace("<", "&lt;")
                   .replace(">", "&gt;")
                   .replace('"', "&quot;")
                   .replace("'", "&#39;"))

    # --------------------------------------------------------------------------
    # UTILITY METHODS
    # --------------------------------------------------------------------------
    def _preview_fallback(self, file_data: bytes, fmt: str, err: str) -> Dict[str, Any]:
        """Fallback preview when extraction fails."""
        return {
            "success": False,
            "error": err,
            "original_size": len(file_data),
            "page_count": 0,
            "estimated_size": 0,
            "estimated_time": 0,
            "complexity": "unknown",
            "recommendations": [],
            "supported_formats": self.supported_formats,
        }

    @staticmethod
    def _estimate_output_size(content: Dict[str, Any], fmt: str) -> int:
        """Estimate output file size."""
        base = len(content.get("text", ""))
        multiplier = {
            "docx": 1.5, 
            "txt": 1.0, 
            "html": 2.0, 
            "images": 10.0,
            "xlsx": 1.3
        }.get(fmt, 1.0)
        return int(base * multiplier)

    @staticmethod
    def _estimate_conversion_time(content: Dict[str, Any], fmt: str) -> int:
        """Estimate conversion time in seconds."""
        pages = content.get("page_count", 1)
        time_per_page = {
            "docx": 2, 
            "txt": 0.5, 
            "html": 1, 
            "images": 3,
            "xlsx": 1.5
        }.get(fmt, 1)
        return max(1, int(pages * time_per_page))

    def _assess_complexity(self, content: Dict[str, Any], fmt: str) -> str:
        """Assess conversion complexity."""
        pages = content.get("page_count", 0)
        tables = len(content.get("tables", []))
        images = len(content.get("images", []))
        
        if pages > 50 or tables > 20 or images > 100:
            return "high"
        if pages > 20 or tables > 10 or images > 50:
            return "medium"
        return "low"

    @staticmethod
    def _get_recommendations(content: Dict[str, Any], fmt: str, opts: Dict[str, Any]) -> List[str]:
        """Get conversion recommendations."""
        rec = []
        pages = content.get("page_count", 0)
        tables = len(content.get("tables", []))
        images = len(content.get("images", []))
        
        if pages > 100:
            rec.append("Large document – consider splitting into smaller parts")
        if tables > 0 and fmt == "txt":
            rec.append("Tables detected – DOCX or XLSX format will preserve formatting better")
        if images > 0 and fmt in ["txt", "xlsx"]:
            rec.append(f"Images detected – they will be lost in {fmt.upper()} format")
        if fmt == "images" and pages > 20:
            rec.append("Many pages – image conversion may take longer and use more storage")
            
        return rec

    @staticmethod
    def _secure_filename(name: str) -> str:
        """Sanitize filename."""
        if not name:
            return "document"
        # Remove unsafe characters
        name = re.sub(r"[^\w\s.-]", "", name)
        name = re.sub(r"[-\s]+", "-", name)
        return name.strip("-")[:100] or "document"

    # --------------------------------------------------------------------------
    # SERVICE STATUS
    # --------------------------------------------------------------------------
    def get_service_status(self) -> Dict[str, Any]:
        """Get service status."""
        return {
            'service_name': 'ConversionService',
            'status': 'operational',
            'supported_formats': self.supported_formats,
            'libraries_available': {
                'pdf': PDF_LIBS_AVAILABLE,
                'docx': DOCX_AVAILABLE,
                'excel': OPENPYXL_AVAILABLE,
                'images': PIL_AVAILABLE
            },
            'file_service_available': self.file_service is not None
        }

    def health_check(self) -> Dict[str, Any]:
        """Perform health check."""
        health_status = {
            'status': 'healthy',
            'service': 'ConversionService',
            'checks': {},
            'errors': []
        }

        try:
            # Check libraries
            lib_check = {
                'pdf_libs': PDF_LIBS_AVAILABLE,
                'docx_libs': DOCX_AVAILABLE,
                'excel_libs': OPENPYXL_AVAILABLE,
                'image_libs': PIL_AVAILABLE
            }
            
            health_status['checks']['libraries'] = lib_check
            
            # Check for missing critical libraries
            if not PDF_LIBS_AVAILABLE:
                health_status['status'] = 'unhealthy'
                health_status['errors'].append("Critical: PDF library (fitz) not available")
            
            # Check for missing optional libraries
            missing_optional = []
            if not DOCX_AVAILABLE:
                missing_optional.append("docx")
            if not OPENPYXL_AVAILABLE:
                missing_optional.append("xlsx")
            if not PIL_AVAILABLE:
                missing_optional.append("images")
                
            if missing_optional:
                health_status['status'] = 'degraded' if health_status['status'] == 'healthy' else health_status['status']
                health_status['errors'].append(f"Optional libraries missing: {', '.join(missing_optional)}")
                
            # Check file service
            if not self.file_service:
                health_status['status'] = 'unhealthy'
                health_status['errors'].append("File management service not available")
                
        except Exception as e:
            health_status['status'] = 'unhealthy'
            health_status['errors'].append(f"Health check failed: {str(e)}")
            
        return health_status